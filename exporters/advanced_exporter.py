#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Exportador Avanzado para Casos de Prueba
Soporte para Word, Linear y Excel con caracteres especiales
"""

import os
import json
import csv
import pandas as pd
from datetime import datetime
from typing import List, Dict, Any
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

class AdvancedExporter:
    """Exportador avanzado con soporte para Word, Linear y Excel mejorado"""
    
    def __init__(self, output_folder: str = "outputs"):
        self.output_folder = output_folder
        os.makedirs(output_folder, exist_ok=True)
    
    def export_to_word(self, test_cases: List[Dict], project_name: str, user_story: str = "") -> str:
        """
        Exporta casos de prueba a Word con formato profesional
        """
        doc = Document()
        
        # Configurar estilos
        self._setup_document_styles(doc)
        
        # Título del documento
        title = doc.add_heading(f'Casos de Prueba: {project_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Información del proyecto
        if user_story:
            doc.add_heading('Historia de Usuario', level=1)
            doc.add_paragraph(user_story)
        
        # Resumen
        doc.add_heading('Resumen', level=1)
        summary_table = doc.add_table(rows=1, cols=2)
        summary_table.style = 'Table Grid'
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells = summary_table.rows[0].cells
        hdr_cells[0].text = 'Métrica'
        hdr_cells[1].text = 'Valor'
        
        # Agregar filas de resumen
        summary_data = [
            ('Total de Casos', str(len(test_cases))),
            ('Fecha de Generación', datetime.now().strftime('%d/%m/%Y %H:%M')),
            ('Casos Funcionales', str(len([tc for tc in test_cases if tc.get('test_type', '').lower() == 'funcional']))),
            ('Casos de Integración', str(len([tc for tc in test_cases if tc.get('test_type', '').lower() == 'integración']))),
            ('Casos Negativos', str(len([tc for tc in test_cases if tc.get('test_type', '').lower() == 'negativo']))),
            ('Casos Límite', str(len([tc for tc in test_cases if tc.get('test_type', '').lower() == 'caso límite']))),
        ]
        
        for metric, value in summary_data:
            row_cells = summary_table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value
        
        # Casos de prueba
        doc.add_heading('Casos de Prueba Detallados', level=1)
        
        for i, test_case in enumerate(test_cases, 1):
            self._add_test_case_to_doc(doc, test_case, i)
        
        # Guardar documento
        filename = f"casos_prueba_{project_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        filepath = os.path.join(self.output_folder, filename)
        doc.save(filepath)
        
        return filepath
    
    def _setup_document_styles(self, doc: Document):
        """Configura estilos del documento"""
        # Estilo para títulos de casos de prueba
        styles = doc.styles
        
        # Crear estilo personalizado para casos de prueba
        if 'Test Case Title' not in [style.name for style in styles]:
            test_case_style = styles.add_style('Test Case Title', 1)  # 1 = paragraph style
            test_case_style.font.name = 'Arial'
            test_case_style.font.size = Pt(14)
            test_case_style.font.bold = True
            test_case_style.font.color.rgb = None  # Color por defecto
    
    def _add_test_case_to_doc(self, doc: Document, test_case: Dict, case_number: int):
        """Agrega un caso de prueba al documento"""
        
        # Título del caso de prueba
        title = doc.add_heading(f'TC-{case_number:03d}: {test_case.get("title", "Sin título")}', level=2)
        
        # Información básica en tabla
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Configurar celdas
        info_data = [
            ('ID:', test_case.get('id', 'N/A')),
            ('Tipo:', test_case.get('test_type', 'N/A')),
            ('Prioridad:', test_case.get('priority', 'N/A')),
            ('Tags:', ', '.join(test_case.get('tags', [])))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = str(value)
        
        # Descripción
        if test_case.get('description'):
            doc.add_heading('Descripción', level=3)
            doc.add_paragraph(test_case['description'])
        
        # Precondiciones
        if test_case.get('preconditions'):
            doc.add_heading('Precondiciones', level=3)
            preconditions = test_case['preconditions']
            if isinstance(preconditions, list):
                for i, precond in enumerate(preconditions, 1):
                    doc.add_paragraph(f'{i}. {precond}', style='List Number')
            else:
                doc.add_paragraph(preconditions)
        
        # Pasos
        if test_case.get('steps'):
            doc.add_heading('Pasos de Ejecución', level=3)
            steps = test_case['steps']
            if isinstance(steps, list):
                for i, step in enumerate(steps, 1):
                    doc.add_paragraph(f'{i}. {step}', style='List Number')
            else:
                doc.add_paragraph(steps)
        
        # Resultado esperado
        if test_case.get('expected_result'):
            doc.add_heading('Resultado Esperado', level=3)
            doc.add_paragraph(test_case['expected_result'])
        
        # Separador
        doc.add_paragraph('─' * 50)
    
    def export_to_linear_enhanced(self, test_cases: List[Dict], project_name: str) -> str:
        """
        Exporta casos de prueba optimizados para Linear
        """
        linear_issues = []
        
        for test_case in test_cases:
            linear_issue = {
                'title': f"TC-{test_case.get('id', '001')}: {test_case.get('title', 'Sin título')}",
                'description': self._create_linear_description(test_case),
                'labels': self._get_linear_labels(test_case),
                'priority': self._map_priority_to_linear(test_case.get('priority', 'medium')),
                'type': 'Test Case',
                'state': 'Todo',
                'assignee': '',  # Se puede asignar manualmente en Linear
                'project': project_name,
                'created_at': datetime.now().isoformat()
            }
            linear_issues.append(linear_issue)
        
        # Exportar a JSON para Linear
        filename = f"linear_import_{project_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.json"
        filepath = os.path.join(self.output_folder, filename)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(linear_issues, f, ensure_ascii=False, indent=2)
        
        # También crear CSV para importación manual
        csv_filename = f"linear_import_{project_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.csv"
        csv_filepath = os.path.join(self.output_folder, csv_filename)
        
        with open(csv_filepath, 'w', newline='', encoding='utf-8') as f:
            if linear_issues:
                writer = csv.DictWriter(f, fieldnames=linear_issues[0].keys())
                writer.writeheader()
                writer.writerows(linear_issues)
        
        return filepath, csv_filepath
    
    def _create_linear_description(self, test_case: Dict) -> str:
        """Crea descripción optimizada para Linear"""
        description = f"**Objetivo:** {test_case.get('title', 'Sin título')}\n\n"
        
        if test_case.get('description'):
            description += f"**Descripción:** {test_case['description']}\n\n"
        
        if test_case.get('preconditions'):
            description += "**Precondiciones:**\n"
            preconditions = test_case['preconditions']
            if isinstance(preconditions, list):
                for i, precond in enumerate(preconditions, 1):
                    description += f"{i}. {precond}\n"
            else:
                description += f"• {preconditions}\n"
            description += "\n"
        
        if test_case.get('steps'):
            description += "**Pasos de Ejecución:**\n"
            steps = test_case['steps']
            if isinstance(steps, list):
                for i, step in enumerate(steps, 1):
                    description += f"{i}. {step}\n"
            else:
                description += f"• {steps}\n"
            description += "\n"
        
        if test_case.get('expected_result'):
            description += f"**Resultado Esperado:** {test_case['expected_result']}\n\n"
        
        # Agregar información técnica
        description += "---\n"
        description += f"**Tipo:** {test_case.get('test_type', 'N/A')}\n"
        description += f"**Prioridad:** {test_case.get('priority', 'N/A')}\n"
        if test_case.get('tags'):
            description += f"**Tags:** {', '.join(test_case['tags'])}\n"
        
        return description
    
    def _get_linear_labels(self, test_case: Dict) -> List[str]:
        """Obtiene etiquetas para Linear"""
        labels = ['Test_Case']
        
        # Agregar tipo como etiqueta
        test_type = test_case.get('test_type', '').lower()
        if test_type:
            labels.append(f"Type_{test_type.replace(' ', '_')}")
        
        # Agregar prioridad como etiqueta
        priority = test_case.get('priority', '').lower()
        if priority:
            labels.append(f"Priority_{priority}")
        
        # Agregar tags personalizados
        if test_case.get('tags'):
            labels.extend([tag.replace(' ', '_') for tag in test_case['tags']])
        
        return labels
    
    def _map_priority_to_linear(self, priority: str) -> str:
        """Mapea prioridad a formato Linear"""
        priority_map = {
            'alta': 'Urgent',
            'high': 'Urgent',
            'media': 'High',
            'medium': 'High',
            'baja': 'Normal',
            'low': 'Normal'
        }
        return priority_map.get(priority.lower(), 'Normal')
    
    def export_to_excel_enhanced(self, test_cases: List[Dict], project_name: str) -> str:
        """
        Exporta a Excel con mejor formato y soporte para caracteres especiales
        """
        # Preparar datos
        data = []
        for tc in test_cases:
            row = {
                'ID': tc.get('id', ''),
                'Título': tc.get('title', ''),
                'Descripción': tc.get('description', ''),
                'Tipo': tc.get('test_type', ''),
                'Prioridad': tc.get('priority', ''),
                'Precondiciones': '\n'.join(tc.get('preconditions', [])) if isinstance(tc.get('preconditions'), list) else tc.get('preconditions', ''),
                'Pasos': '\n'.join([f"{i+1}. {step}" for i, step in enumerate(tc.get('steps', []))]) if isinstance(tc.get('steps'), list) else tc.get('steps', ''),
                'Resultado Esperado': tc.get('expected_result', ''),
                'Tags': ', '.join(tc.get('tags', [])),
                'Historia de Usuario': tc.get('user_story', '')
            }
            data.append(row)
        
        # Crear DataFrame
        df = pd.DataFrame(data)
        
        # Guardar con encoding UTF-8
        filename = f"casos_prueba_{project_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
        filepath = os.path.join(self.output_folder, filename)
        
        with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Casos de Prueba', index=False)
            
            # Obtener la hoja de trabajo
            worksheet = writer.sheets['Casos de Prueba']
            
            # Ajustar ancho de columnas
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                worksheet.column_dimensions[column_letter].width = adjusted_width
        
        return filepath

def main():
    """Función principal para pruebas"""
    # Ejemplo de uso
    test_cases = [
        {
            'id': 'TC-001',
            'title': 'Verificar login con credenciales válidas',
            'description': 'Verificar que el usuario puede iniciar sesión con credenciales correctas',
            'test_type': 'Funcional',
            'priority': 'Alta',
            'preconditions': ['Usuario registrado en el sistema', 'Credenciales válidas disponibles'],
            'steps': ['Navegar a la página de login', 'Ingresar email válido', 'Ingresar contraseña válida', 'Hacer clic en "Iniciar Sesión"'],
            'expected_result': 'El usuario es redirigido al dashboard principal',
            'tags': ['login', 'autenticación', 'happy-path'],
            'user_story': 'Como usuario del sistema, quiero poder iniciar sesión'
        }
    ]
    
    exporter = AdvancedExporter()
    
    # Exportar a Word
    word_file = exporter.export_to_word(test_cases, "Sistema de Login", "Historia de usuario de login")
    print(f"✅ Exportado a Word: {word_file}")
    
    # Exportar a Linear
    linear_json, linear_csv = exporter.export_to_linear_enhanced(test_cases, "Sistema de Login")
    print(f"✅ Exportado a Linear JSON: {linear_json}")
    print(f"✅ Exportado a Linear CSV: {linear_csv}")
    
    # Exportar a Excel mejorado
    excel_file = exporter.export_to_excel_enhanced(test_cases, "Sistema de Login")
    print(f"✅ Exportado a Excel: {excel_file}")

if __name__ == "__main__":
    main()
